"""Tests for the file handling / text extraction module."""

import os
import pytest
import tempfile
from app.file_handling.parser import (
    extract_text,
    extract_text_from_txt,
)


def test_extract_text_from_txt():
    """Test plain text file extraction."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("Hallo Welt!\nDies ist ein Test.")
        f.flush()
        temp_path = f.name

    try:
        result = extract_text_from_txt(temp_path)
        assert "Hallo Welt!" in result
        assert "Dies ist ein Test." in result
    finally:
        os.unlink(temp_path)


def test_extract_text_dispatcher_txt():
    """Test that extract_text dispatches .txt files correctly."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("Testinhalt")
        f.flush()
        temp_path = f.name

    try:
        result = extract_text(temp_path)
        assert "Testinhalt" in result
    finally:
        os.unlink(temp_path)


def test_extract_text_unsupported_format():
    """Test that unsupported file formats raise ValueError."""
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"test")
        temp_path = f.name

    try:
        with pytest.raises(ValueError, match="Unsupported file format"):
            extract_text(temp_path)
    finally:
        os.unlink(temp_path)


def test_extract_text_from_txt_unicode():
    """Test that unicode text is extracted correctly."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("Ünïcödë Tëxt mit Ümläutën: äöüß")
        f.flush()
        temp_path = f.name

    try:
        result = extract_text_from_txt(temp_path)
        assert "Ümläutën" in result
        assert "äöüß" in result
    finally:
        os.unlink(temp_path)


def test_extract_text_from_docx():
    """Test DOCX extraction with a real docx file."""
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = f.name

    try:
        doc = Document()
        doc.add_paragraph("Erster Absatz")
        doc.add_paragraph("Zweiter Absatz")
        doc.save(temp_path)

        result = extract_text(temp_path)
        assert "Erster Absatz" in result
        assert "Zweiter Absatz" in result
    finally:
        os.unlink(temp_path)
